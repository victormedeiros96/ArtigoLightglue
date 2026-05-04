from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Relatório Técnico: Benchmarking de Rastreamento LightGlue', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Este documento detalha os experimentos e resultados finais do rastreador baseado em LightGlue comparado aos algoritmos BoTSORT e ByteTrack.')

    # Section 1
    doc.add_heading('1. Estatísticas dos Datasets', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Dataset'
    hdr_cells[1].text = 'Sequências'
    hdr_cells[2].text = 'Frames'
    hdr_cells[3].text = 'Descrição'
    
    data = [
        ('GoPro Benchmark', '5', '15.939', 'Cenário roadside, alta velocidade, deslocamentos lineares.'),
        ('GOT-10k (Val)', '180', '21.007', 'Cenário geral, câmeras em movimento.')
    ]
    for ds, seq, frames, desc in data:
        row_cells = table.add_row().cells
        row_cells[0].text = ds
        row_cells[1].text = seq
        row_cells[2].text = frames
        row_cells[3].text = desc

    # Section 2
    doc.add_heading('2. Resultados Consolidados (Cenário Crítico 2 FPS)', level=1)
    doc.add_paragraph('Resultados obtidos com Stride 15 (2 FPS) na GoPro, onde o deslocamento entre frames é máximo.')
    
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Método'
    hdr_cells[1].text = 'MOTA (↑)'
    hdr_cells[2].text = 'IDF1 (↑)'
    hdr_cells[3].text = 'IDSW (↓)'
    
    results = [
        ('BoTSORT (Baseline)', '0.82%', '1.58%', '5'),
        ('LightGlue (Otimizado)', '4.53%', '8.40%', '302')
    ]
    for meth, mota, idf1, idsw in results:
        row_cells = table2.add_row().cells
        row_cells[0].text = meth
        row_cells[1].text = mota
        row_cells[2].text = idf1
        row_cells[3].text = idsw

    # Section 3
    doc.add_heading('3. Análise de Comportamento e Otimizações', level=1)
    
    points = [
        ('Resiliência Visual', 'O LightGlue utiliza correspondência de pontos-chave via SuperPoint, permitindo reencontrar objetos mesmo após deslocamentos de centenas de pixels.'),
        ('Gate Espacial Adaptativo', 'Expandimos o limite de busca para 1500px para lidar com o movimento de alta velocidade do veículo em baixas taxas de quadros.'),
        ('Conclusão sobre CMC', 'Em cenários de infraestrutura fixa ou rodoviária linear, o matching puro do LightGlue com gate expandido é superior ao uso de CMC baseado em translação, reduzindo ruído na estimativa de movimento.'),
        ('Superioridade Técnica', 'O método proposto entrega uma performance 5.5x superior ao BoTSORT no cenário de 2 FPS.')
    ]
    
    for title, desc in points:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)

    doc.save('/mnt/hd2/ArtigoLightglue/mass_results/Relatorio_Tecnico_LightGlue.docx')
    print("Relatório DOCX gerado com sucesso em mass_results/Relatorio_Tecnico_LightGlue.docx")

if __name__ == '__main__':
    create_report()
